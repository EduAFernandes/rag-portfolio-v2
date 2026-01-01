"""
Data loading module for ingesting DOCX files from MinIO.

This module provides functionality to connect to MinIO (S3-compatible storage)
and load DOCX documents with proper parsing and initial processing.
"""

from __future__ import annotations

import hashlib
import io
import tempfile
from collections.abc import Generator
from datetime import datetime
from pathlib import Path
from typing import TYPE_CHECKING, Any, Optional

from minio import Minio

if TYPE_CHECKING:
    from typing import Dict
from minio.error import S3Error
from llama_index.core import Document
from llama_index.core.readers import SimpleDirectoryReader
from docx import Document as DocxDocument
import pandas as pd
from loguru import logger
from tenacity import retry, stop_after_attempt, wait_exponential

from config import MinIOConfig, Settings


class MinIODocumentLoader:
    """
    Loader for DOCX documents stored in MinIO.
    
    This class handles connection to MinIO, document listing, downloading,
    and initial parsing of DOCX files into LlamaIndex Document objects.
    """
    
    def __init__(self, settings: Settings) -> None:
        """
        Initialize MinIO document loader.
        
        Args:
            settings: Application settings containing MinIO configuration
        """
        self.settings = settings
        self.minio_config = settings.minio
        self.client = self._initialize_client()
        self._ensure_bucket_exists()
        
    def _initialize_client(self) -> Minio:
        """
        Initialize MinIO client with configuration.
        
        Returns:
            Minio: Configured MinIO client
        """
        endpoint = self.minio_config.endpoint.replace("http://", "").replace("https://", "")
        
        client = Minio(
            endpoint=endpoint,
            access_key=self.minio_config.access_key,
            secret_key=self.minio_config.secret_key,
            secure=self.minio_config.secure,
            region=self.minio_config.region
        )
        
        logger.info(f"MinIO client initialized for endpoint: {endpoint}")
        return client
    
    def _ensure_bucket_exists(self) -> None:
        """Create bucket if it doesn't exist."""
        try:
            if not self.client.bucket_exists(self.minio_config.bucket_name):
                self.client.make_bucket(
                    self.minio_config.bucket_name,
                    location=self.minio_config.region
                )
                logger.info(f"Created bucket: {self.minio_config.bucket_name}")
            else:
                logger.info(f"Bucket exists: {self.minio_config.bucket_name}")
        except S3Error as e:
            logger.error(f"Error checking/creating bucket: {e}")
            raise
    
    @retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=4, max=10))
    def list_docx_files(self) -> list[Dict[str, Any]]:
        """
        List all DOCX files in the configured bucket.
        
        Returns:
            List of dictionaries containing file metadata
        """
        files = []
        
        try:
            objects = self.client.list_objects(
                self.minio_config.bucket_name,
                recursive=True
            )
            
            for obj in objects:
                if obj.object_name.lower().endswith('.docx'):
                    files.append({
                        'name': obj.object_name,
                        'size': obj.size,
                        'last_modified': obj.last_modified,
                        'etag': obj.etag,
                        'content_type': obj.content_type or 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                    })
                    
            logger.info(f"Found {len(files)} DOCX files in bucket")
            return files
            
        except S3Error as e:
            logger.error(f"Error listing files: {e}")
            raise
    
    @retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=4, max=10))
    def download_file(self, object_name: str) -> bytes:
        """
        Download a file from MinIO.
        
        Args:
            object_name: Name/path of the object in MinIO
            
        Returns:
            File content as bytes
        """
        try:
            response = self.client.get_object(
                self.minio_config.bucket_name,
                object_name
            )
            data = response.read()
            response.close()
            response.release_conn()
            
            logger.debug(f"Downloaded {object_name}: {len(data)} bytes")
            return data
            
        except S3Error as e:
            logger.error(f"Error downloading {object_name}: {e}")
            raise
    
    def parse_docx_content(self, content: bytes, filename: str) -> Dict[str, Any]:
        """
        Parse DOCX content and extract structured data.
        
        Args:
            content: DOCX file content as bytes
            filename: Original filename for metadata
            
        Returns:
            Dictionary containing parsed content and metadata
        """
        doc_data = {
            'filename': filename,
            'content': [],
            'tables': [],
            'metadata': {},
            'structure': []
        }
        
        try:
            doc = DocxDocument(io.BytesIO(content))
            
            # Extract metadata
            doc_data['metadata'] = {
                'created': doc.core_properties.created,
                'modified': doc.core_properties.modified,
                'author': doc.core_properties.author,
                'title': doc.core_properties.title,
                'subject': doc.core_properties.subject,
                'keywords': doc.core_properties.keywords,
                'category': doc.core_properties.category,
                'content_hash': hashlib.sha256(content).hexdigest()
            }
            
            # Extract paragraphs with styling information
            for para_idx, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    para_data = {
                        'index': para_idx,
                        'text': paragraph.text,
                        'style': paragraph.style.name if paragraph.style else None,
                        'level': self._get_heading_level(paragraph),
                        'alignment': str(paragraph.alignment) if paragraph.alignment else None
                    }
                    doc_data['content'].append(para_data)
                    doc_data['structure'].append(('paragraph', para_idx, para_data['level']))
            
            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_data = self._extract_table_data(table)
                if table_data:
                    doc_data['tables'].append({
                        'index': table_idx,
                        'data': table_data,
                        'rows': len(table.rows),
                        'columns': len(table.columns)
                    })
                    doc_data['structure'].append(('table', table_idx, None))
            
            logger.info(f"Parsed {filename}: {len(doc_data['content'])} paragraphs, {len(doc_data['tables'])} tables")
            
        except Exception as e:
            logger.error(f"Error parsing DOCX content for {filename}: {e}")
            raise
        
        return doc_data
    
    def _get_heading_level(self, paragraph) -> Optional[int]:
        """
        Determine heading level from paragraph style.
        
        Args:
            paragraph: python-docx paragraph object
            
        Returns:
            Heading level (1-9) or None if not a heading
        """
        if paragraph.style and paragraph.style.name:
            style_name = paragraph.style.name.lower()
            if 'heading' in style_name:
                try:
                    level = int(style_name.split()[-1])
                    return level
                except (ValueError, IndexError):
                    return 0
        return None
    
    def _extract_table_data(self, table) -> list[list[str]]:
        """
        Extract data from a DOCX table.
        
        Args:
            table: python-docx table object
            
        Returns:
            List of rows, each containing list of cell values
        """
        table_data = []
        
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_data.append(cell_text)
            if any(row_data):
                table_data.append(row_data)
        
        return table_data
    
    def create_llamaindex_documents(self, parsed_data: Dict[str, Any]) -> list[Document]:
        """
        Convert parsed DOCX data to LlamaIndex Document objects.
        
        Args:
            parsed_data: Parsed document data from parse_docx_content
            
        Returns:
            List of LlamaIndex Document objects
        """
        documents = []
        
        # Combine content into structured text
        full_text_parts = []
        
        for item_type, item_idx, _ in parsed_data['structure']:
            if item_type == 'paragraph':
                if item_idx < len(parsed_data['content']):
                    para = parsed_data['content'][item_idx]
                    full_text_parts.append(para['text'])
            elif item_type == 'table':
                # Find the table with matching index
                table_found = None
                for table in parsed_data['tables']:
                    if table.get('index') == item_idx:
                        table_found = table
                        break
                if table_found:
                    table_text = self._format_table_as_text(table_found['data'])
                    full_text_parts.append(f"\n[Table {item_idx + 1}]\n{table_text}\n")
        
        full_text = "\n\n".join(full_text_parts)
        
        # Extract module information from filename
        module_info = self._extract_module_info(parsed_data['filename'])
        
        # Create main document
        doc = Document(
            text=full_text,
            metadata={
                **parsed_data['metadata'],
                'filename': parsed_data['filename'],
                'module': module_info.get('module', 'unknown'),
                'module_number': module_info.get('number', 0),
                'document_type': 'docx',
                'paragraph_count': len(parsed_data['content']),
                'table_count': len(parsed_data['tables']),
                'processing_timestamp': datetime.utcnow().isoformat(),
                'source': f"minio://{self.minio_config.bucket_name}/{parsed_data['filename']}"
            }
        )
        
        documents.append(doc)
        
        # Optionally create separate documents for tables
        for table_info in parsed_data['tables']:
            table_doc = Document(
                text=self._format_table_as_text(table_info['data']),
                metadata={
                    'filename': parsed_data['filename'],
                    'module': module_info.get('module', 'unknown'),
                    'content_type': 'table',
                    'table_index': table_info['index'],
                    'rows': table_info['rows'],
                    'columns': table_info['columns'],
                    'source': f"minio://{self.minio_config.bucket_name}/{parsed_data['filename']}"
                }
            )
            documents.append(table_doc)
        
        return documents
    
    def _extract_module_info(self, filename: str) -> Dict[str, Any]:
        """
        Extract module information from filename.
        
        Args:
            filename: Name of the file (e.g., 'mod-1-al-1.docx')
            
        Returns:
            Dictionary with module information
        """
        import re
        
        pattern = r'mod-(\d+)-.*'
        match = re.match(pattern, filename.lower())
        
        if match:
            module_num = int(match.group(1))
            return {
                'module': f'mod-{module_num}',
                'number': module_num
            }
        
        return {'module': 'unknown', 'number': 0}
    
    def _format_table_as_text(self, table_data: List[List[str]]) -> str:
        """
        Format table data as readable text.
        
        Args:
            table_data: List of table rows
            
        Returns:
            Formatted table as text
        """
        if not table_data:
            return ""
        
        try:
            if len(table_data) > 1:
                df = pd.DataFrame(table_data[1:], columns=table_data[0])
            else:
                df = pd.DataFrame(table_data)
            return df.to_string(index=False)
        except Exception as e:
            # Fallback to simple text formatting
            return "\n".join([" | ".join(row) for row in table_data])
    
    def load_all_documents(self) -> Generator[Document, None, None]:
        """
        Load all DOCX documents from MinIO as LlamaIndex Documents.
        
        Yields:
            Document objects one at a time for memory efficiency
        """
        files = self.list_docx_files()
        
        for file_info in files:
            try:
                logger.info(f"Processing {file_info['name']}")
                
                # Download file
                content = self.download_file(file_info['name'])
                
                # Parse DOCX
                parsed_data = self.parse_docx_content(content, file_info['name'])
                
                # Convert to LlamaIndex documents
                documents = self.create_llamaindex_documents(parsed_data)
                
                for doc in documents:
                    yield doc
                    
            except Exception as e:
                logger.error(f"Failed to process {file_info['name']}: {e}")
                continue
    
    def load_specific_document(self, object_name: str) -> list[Document]:
        """
        Load a specific DOCX document from MinIO.
        
        Args:
            object_name: Name/path of the document in MinIO
            
        Returns:
            List of Document objects
        """
        try:
            content = self.download_file(object_name)
            parsed_data = self.parse_docx_content(content, object_name)
            return self.create_llamaindex_documents(parsed_data)
        except Exception as e:
            logger.error(f"Failed to load {object_name}: {e}")
            raise